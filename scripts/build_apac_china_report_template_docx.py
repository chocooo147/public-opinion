#!/usr/bin/env python3
"""Build the China-only APAC weekly community sentiment Word template."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "templates" / "APAC_Weekly_Community_Sentiment_Report_Template.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
FONT = "Calibri"
INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
APEX_RED = "C51B2A"  # Named brand override for the kicker and title rule only.


def set_run_font(run, *, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == CONTENT_WIDTH_DXA, widths
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def format_cell(cell, text, *, header=False, center=False, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=INK, bold=header)
    if header:
        shade_cell(cell, LIGHT_FILL)


def add_table(doc, headers, rows, widths, *, centers=None, font_size=9, trailing_space=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    centers = centers or set()
    for idx, value in enumerate(headers):
        format_cell(table.rows[0].cells[idx], value, header=True, center=idx in centers, size=font_size)
    repeat_header(table.rows[0])
    keep_row_together(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        keep_row_together(row)
        for idx, value in enumerate(row_data):
            format_cell(row.cells[idx], value, center=idx in centers, size=font_size)
    set_table_geometry(table, widths)
    if trailing_space:
        doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_para(doc, text="", *, size=11, color=INK, bold=False, italic=False, align=None, before=0, after=6):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_callout(doc, label, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.10
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.keep_together = True
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "start", "bottom", "end"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "7")
        border.set(qn("w:color"), "D7DBE2")
        p_bdr.append(border)
    p_pr.append(p_bdr)
    run = paragraph.add_run(f"{label}: ")
    set_run_font(run, size=10, color=DARK_BLUE, bold=True)
    run = paragraph.add_run(text)
    set_run_font(run, size=10, color=INK)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_title_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), APEX_RED)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_section_heading(doc, number, title):
    return doc.add_heading(f"{number}. {title}", level=1)


def build_document():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("APAC COMMUNITY SENTIMENT  /  CHINA")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("CHINA WEEKLY REPORT TEMPLATE  •  PAGE ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(footer, "PAGE")

    add_para(doc, "APEX REPORTING TEMPLATE", size=10, color=APEX_RED, bold=True, after=6)
    title = add_para(doc, "APAC — CHINA WEEKLY\nCOMMUNITY SENTIMENT REPORT", size=24, color=INK, bold=True, after=14)
    title.paragraph_format.line_spacing = 1.0
    add_title_rule(title)
    add_para(doc, "A China-only executive template for evidence-led weekly community reporting", size=12, color=MUTED, after=18)

    add_table(
        doc,
        ["REPORT FIELD", "VALUE", "REPORT FIELD", "VALUE"],
        [
            ["Reporting period", "[YYYY/MM/DD–YYYY/MM/DD]", "Region", "China"],
            ["Game / product", "[Name]", "Prepared by", "[Team / Analyst]"],
            ["Data sources", "[China platforms / tools]", "Template version", "3.0"],
        ],
        [1500, 3180, 1500, 3180],
        centers={0, 2},
        font_size=9.5,
    )
    add_callout(
        doc,
        "SCOPE",
        "China only. The China reporting team does not aggregate other regional reports. Use a complete natural week and label missing, estimated, or sample-limited inputs explicitly.",
    )
    add_callout(
        doc,
        "NARRATIVE SEQUENCE",
        "Conclusion → Attribution → Evidence → Comparison → Impact. Present drivers in the order Positive → Neutral or Mixed → Negative.",
    )

    add_section_heading(doc, 1, "OVERVIEW")
    doc.add_heading("Executive Summary", level=2)
    add_callout(doc, "WRITE 80–140 WORDS", "State the China community direction first, identify the most material opportunity and risk, explain the main week-over-week change, and end with the expected impact.")
    add_para(doc, "[Executive summary]", color=MUTED, italic=True, after=10)
    doc.add_heading("China Weekly Snapshot", level=2)
    add_table(
        doc,
        ["METRIC", "CURRENT", "PREVIOUS", "WEEK-OVER-WEEK"],
        [
            ["Overall sentiment metric [define]", "[ ]", "[ ]", "[↑ / ↓ / =  change]"],
            ["Primary positive driver", "[Topic]", "[Topic]", "[What changed]"],
            ["Primary neutral / mixed driver", "[Topic]", "[Topic]", "[What changed]"],
            ["Primary negative driver", "[Topic]", "[Topic]", "[What changed]"],
        ],
        [3000, 1800, 1800, 2760],
        centers={1, 2},
        font_size=9.3,
    )
    add_para(doc, "China summary — 120–180 Chinese characters or 60–100 English words", size=10, color=DARK_BLUE, bold=True, after=4)
    add_para(doc, "[Conclusion, key drivers, observable evidence, previous-week comparison, and likely impact.]", color=MUTED, italic=True)

    add_section_heading(doc, 2, "TOP CONVERSATION DRIVERS")
    add_para(doc, "List only material China drivers. If fewer than ten meet the reporting threshold, do not call this section “Top 10.”", size=10, color=MUTED)
    add_table(
        doc,
        ["#", "DRIVER", "CURRENT", "PREVIOUS", "WOW", "LABEL", "EVIDENCE-LED NARRATIVE"],
        [
            ["1", "[Topic]", "[X.X% / N]", "[X.X% / N]", "[New / ↑ / ↓ / =]", "[Positive / Neutral / Mixed / Negative]", "[Reaction + reason + comparison + impact]"],
            ["2", "[Topic]", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"],
            ["3", "[Topic]", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"],
            ["4", "[Topic]", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"],
            ["5", "[Topic]", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"],
        ],
        [520, 1500, 1080, 1080, 900, 1500, 2780],
        centers={0, 2, 3, 4, 5},
        font_size=8.2,
    )
    doc.add_heading("China Topic Notes", level=2)
    add_para(doc, "[Topic] — [Positive / Neutral / Mixed / Negative] — [Storyline stage]", size=10.5, color=DARK_BLUE, bold=True, after=4)
    add_para(doc, "Players [reaction or behavior] because [reason]. Compared with [previous week / expectation / earlier version], [difference]. This contributed to [sentiment / participation / viewership / trust].", color=MUTED, italic=True)
    add_callout(doc, "MIXED-SENTIMENT PATTERN", "Players welcomed [positive aspect], but concerns remained around [negative aspect]. Compared with [reference point], the feature was seen as [relative judgment], resulting in an overall mixed response.")
    add_para(doc, "For recurring topics, state what changed and advance the storyline: anticipation → launch → initial reaction → evaluation → problem attribution → clarification → resolution or persistence.", size=10, color=MUTED)

    add_section_heading(doc, 3, "SENTIMENT HISTORY")
    doc.add_heading("Weekly Trend", level=2)
    add_table(
        doc,
        ["WEEK", "CHINA METRIC", "WOW", "POSITIVE DRIVER", "NEUTRAL / MIXED", "NEGATIVE DRIVER", "STAGE"],
        [
            ["[W-3]", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"],
            ["[W-2]", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"],
            ["[W-1]", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"],
            ["[Current]", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"],
        ],
        [880, 1200, 900, 1600, 1600, 1600, 1580],
        centers={0, 1, 2},
        font_size=8.4,
    )
    doc.add_heading("Trend Interpretation", level=2)
    add_para(doc, "Over the past [number] weeks, China community sentiment has [trend], primarily influenced by [recurring topic]. The current week marks the [storyline stage]. Compared with [reference point], this appears to have affected [trust / participation / feature adoption / viewership].", color=MUTED, italic=True)

    add_section_heading(doc, 4, "UGC / STREAM VIEWERSHIP")
    add_callout(doc, "INTERPRETATION RULE", "Metric change → event or content driver → meaning. Separate reach, depth, intensity, and creator activity. Never substitute discussion volume for a missing official viewing metric.")
    add_table(
        doc,
        ["METRIC", "CURRENT", "PREVIOUS", "WOW", "PRIMARY DRIVER", "INTERPRETATION"],
        [
            ["Reach — Total views", "[ ]", "[ ]", "[ ]", "[ ]", "[Broader / narrower reach]"],
            ["Depth — Watch hours", "[ ]", "[ ]", "[ ]", "[ ]", "[Depth of viewing engagement]"],
            ["Intensity — Popularity / heat", "[ ]", "[ ]", "[ ]", "[ ]", "[Viewing / discussion intensity]"],
            ["Creator activity — UGC volume", "[ ]", "[ ]", "[ ]", "[ ]", "[Content supply]"],
            ["Creator activity — Active creators", "[ ]", "[ ]", "[ ]", "[ ]", "[Breadth of participation]"],
        ],
        [1750, 950, 950, 850, 1900, 2960],
        centers={1, 2, 3},
        font_size=8.7,
    )
    doc.add_heading("China Viewership Interpretation", level=2)
    add_para(doc, "[Metric] changed from [previous value] to [current value] ([WoW change]), primarily driven by [event or content]. Because [metric definition], interpret it as [reach / depth / intensity / creator activity]. State the likely impact without claiming unsupported causality.", color=MUTED, italic=True)

    add_section_heading(doc, 5, "METHODOLOGY NOTES")
    methodology = [
        ("Reporting scope", "China only; no aggregation of other regional reports."),
        ("Natural-week rule", "[Monday 00:00 through Sunday 23:59; timezone]"),
        ("Sentiment scale and thresholds", "[Fixed system; define Positive, Neutral, Mixed, Negative]"),
        ("Topic model", "[Single-label / multi-label; model and version]"),
        ("China data sources and coverage", "[Platforms, content types, dates, exclusions]"),
        ("Platform-unit comparability", "[Views, comments, posts, popularity, or heat; incompatible units are not summed]"),
        ("Viewership metric definitions", "[Views / watch hours / popularity / heat / UGC / creators]"),
        ("Methodology changes", "[None / change and impact]"),
        ("Missing data", "[None / estimated / excluded]"),
        ("Causality", "[Association / hypothesis / official confirmation]"),
    ]
    methodology_rows = []
    for index in range(0, len(methodology), 2):
        methodology_rows.append([*methodology[index], *methodology[index + 1]])
    add_table(
        doc,
        ["METHOD FIELD", "DISCLOSURE", "METHOD FIELD", "DISCLOSURE"],
        methodology_rows,
        [1650, 3030, 1650, 3030],
        font_size=8.4,
    )

    add_section_heading(doc, 6, "DATA QUALITY CHECK")
    doc.add_heading("Completion Status", level=2)
    checks = [
        "Reporting scope is China only.",
        "Reporting period is a complete natural week.",
        "All required sections remain in the fixed order.",
        "Overview follows Positive → Neutral or Mixed → Negative.",
        "Current and previous-week values use the same definition.",
        "Percentage changes are mathematically correct and include their basis.",
        "Sentiment thresholds are consistent or the change is disclosed.",
        "Recurring topics state what changed week over week.",
        "Mixed topics use concession logic.",
        "Incompatible platform units are not summed or directly compared.",
        "Missing official metrics are marked missing rather than inferred.",
        "Unverified causes are labeled as hypotheses.",
    ]
    paired_checks = []
    for index in range(0, len(checks), 2):
        paired_checks.append(["☐", checks[index], "☐", checks[index + 1]])
    add_table(
        doc,
        ["", "QUALITY GATE", "", "QUALITY GATE"],
        paired_checks,
        [520, 4160, 520, 4160],
        centers={0, 2},
        font_size=8.8,
    )
    doc.add_heading("Data-Quality Warnings", level=2)
    add_table(
        doc,
        ["#", "WARNING", "AFFECTED METRIC / PLATFORM", "INTERPRETATION IMPACT"],
        [[str(i), "[Warning]", "[ ]", "[ ]"] for i in range(1, 4)],
        [600, 3000, 2400, 3360],
        centers={0},
        font_size=9,
    )
    doc.add_heading("Analyst Notes", level=2)
    add_para(doc, "[State missing, estimated, sample-limited, changed, or non-comparable inputs and whether the report remains a draft.]", size=9.2, color=MUTED, italic=True, after=2)
    doc.add_heading("Next-Week Watchlist", level=2)
    add_table(
        doc,
        ["METRIC / TOPIC", "CURRENT BASELINE", "TRIGGER", "WHY IT MATTERS"],
        [
            ["[Sentiment metric]", "[ ]", "[e.g. ±0.5]", "[Decision impact]"],
            ["[Negative topic share]", "[ ]", "[e.g. >10%]", "[Decision impact]"],
        ],
        [2400, 2000, 1900, 3060],
        centers={1, 2},
        font_size=9.2,
        trailing_space=False,
    )

    properties = doc.core_properties
    properties.title = "APAC — China Weekly Community Sentiment Report Template"
    properties.subject = "China-only weekly community sentiment reporting template"
    properties.author = "APEX Reporting"
    properties.keywords = "APAC, China, weekly report, community sentiment"
    properties.comments = "Template version 3.0"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
